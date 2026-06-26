#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Urban-Renewal 專案狀態報告 → Word (.docx)
外商顧問格式（Cover / Executive Summary / 現況 / 資產盤點 / 技術與品質 / 法規查證 / 發現 / 建議與展望 / 附錄）。
內容同 docs/Urban-Renewal-狀態報告-2026-06.md；法規/市場數據經 2026-06 網路查證。
用法：python3 make_report_docx.py   輸出：Urban-Renewal-狀態報告-2026-06.docx   相依：pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EMER=RGBColor(0x06,0x4E,0x3B); EMERL=RGBColor(0x10,0x98,0x71); INK=RGBColor(0x1C,0x1A,0x17)
BODY=RGBColor(0x44,0x40,0x3C); MUTE=RGBColor(0x8A,0x83,0x7C); WHITE=RGBColor(0xFF,0xFF,0xFF)
RED=RGBColor(0x9A,0x34,0x12); AMB=RGBColor(0xB4,0x53,0x1C)
WARM="F1EEE6"; MINT="E9F7F0"; HEAD="0A5D46"
SERIF="Noto Serif TC"; SANS="Microsoft JhengHei"

def _ea(run,name):
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'),name); rf.set(qn('w:ascii'),name); rf.set(qn('w:hAnsi'),name)
def shade(cell,hexc):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc)
    cell._tc.get_or_add_tcPr().append(sh)
def run(p,t,size=10.5,color=BODY,bold=False,font=SANS):
    r=p.add_run(t); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name=font; _ea(r,font); return r
def para(doc,t="",size=10.5,color=BODY,bold=False,font=SANS,align=None,before=2,after=4,line=1.3):
    p=doc.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line
    if align: p.alignment=align
    if t: run(p,t,size,color,bold,font)
    return p
def bullets(doc,items,size=10.5):
    for it in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.28
        if isinstance(it,tuple): run(p,it[0]+'：',size,INK,bold=True); run(p,it[1],size,BODY)
        else: run(p,it,size,BODY)
def heading(doc,t,level,kicker=None):
    if kicker:
        k=doc.add_paragraph(); k.paragraph_format.space_before=Pt(13); k.paragraph_format.space_after=Pt(0)
        run(k,kicker.upper(),8.5,EMERL,bold=True)
    p=doc.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(2 if kicker else 11); pf.space_after=Pt(5)
    run(p,t,{1:22,2:16,3:12.5}.get(level,12),EMER if level<=2 else INK,bold=True,font=SERIF)
def callout(doc,label,text,kind="emer"):
    fill={"emer":MINT,"warn":"FBF3E6"}.get(kind,MINT); bar={"emer":EMER,"warn":AMB}.get(kind,EMER)
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,fill)
    c.paragraphs[0].paragraph_format.space_after=Pt(1); run(c.paragraphs[0],label,8.5,bar,bold=True)
    p2=c.add_paragraph(); p2.paragraph_format.space_before=Pt(1); run(p2,text,10,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,HEAD); c.paragraphs[0].paragraph_format.space_after=Pt(1); c.paragraphs[0].paragraph_format.space_before=Pt(1)
        run(c.paragraphs[0],h,9,WHITE,bold=True)
    for ri,rd in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(rd):
            if ri%2==1: shade(cells[i],WARM)
            cells[i].paragraphs[0].paragraph_format.space_after=Pt(1); cells[i].paragraphs[0].paragraph_format.space_before=Pt(1)
            run(cells[i].paragraphs[0],v,8.7,BODY)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def build():
    doc=Document(); st=doc.styles['Normal']; st.font.name=SANS; st.font.size=Pt(10.5)
    sec=doc.sections[0]; sec.page_height=Cm(29.7); sec.page_width=Cm(21)
    sec.top_margin=sec.bottom_margin=Cm(2.0); sec.left_margin=sec.right_margin=Cm(2.0)

    # 封面
    para(doc,before=46)
    para(doc,"PROJECT STATUS REPORT · 專案狀態報告",10,EMERL,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    run(p,"Urban-Renewal\n都市更新／危老開發評估數位體系",30,EMER,bold=True,font=SERIF)
    para(doc,"知識庫 · 互動儀表板 · 策略沙盤 · 開發試算 · 對外簡報",13,BODY,bold=True,before=4,after=22)
    table(doc,["項目","內容"],[
        ["報告對象","專案負責人／開發部"],["報告日期","2026-06-22"],
        ["報告範圍","jeremy0819/Urban-Renewal 公開知識庫 + 互動工具組"],
        ["基準提交","5653e22（main，第 18 次提交）"],
        ["查證原則","法規／稅制／市場數據均經 2026-06 網路查證並附來源"],
        ["資料邊界","全庫去識別化，不含真實案名／金額／人名"],
    ],widths=[3.6,13.4])
    doc.add_page_break()

    # 一、執行摘要
    heading(doc,"執行摘要 Executive Summary",2,"01")
    callout(doc,"一句話","本專案已從 3 份方法論文件，成長為一套可上線、零依賴、經事實查證的都更／危老數位決策與訓練體系；核心缺口在遊戲真實感調校與 RE-DCF-Tool 串接。")
    bullets(doc,[
        ("完成度","高。5 個自含式網頁（2,583 行）＋3 份方法論＋白皮書 Word＋說明會 PPT＋5 份 spec，可離線開、可上 Pages。"),
        ("品質與資安","良好。純前端、零外送、無密鑰、無 XSS、外連皆 noopener、.gitignore 鐵律生效、洩漏掃描乾淨。"),
        ("內容正確性","經查證。§37／§65／§67／§162、房地合一 2.0／囤房稅 2.0、BERS、2026 房市趨勢均對照官方與媒體來源。"),
        ("主要缺口","① 沙盤『人最難整合』張力不足；② 競爭建商為展示性；③ 財務單向、未串 RE-DCF-Tool。"),
        ("建議下一步","先做低成本『真實感調校包』，再進 RE-DCF-Tool Phase 0／1 串接，逐步邁向 V4–V7。"),
    ])

    # 二、現況
    heading(doc,"專案現況與定位 Situation",2,"02")
    para(doc,"Urban-Renewal 為方法論主庫（知識層），與姊妹庫 RE-DCF-Tool（計算層，Streamlit）分工。將「拿到一塊地→交屋」拆為 S1–S11 全生命週期，以不動產研析為脊椎。")
    para(doc,"本期經 18 次提交，由純文件演進為四層交付：① 知識層（流程架構／投報正典／整合總論）② 互動層（儀表板／策略沙盤／開發試算）③ 對外層（說明會簡報）④ 工程層（spec／產生器／資安掃描）。")

    # 三、資產盤點
    heading(doc,"資產盤點 Asset Inventory",2,"03")
    table(doc,["交付物","規模","狀態","用途"],[
        ["index.html 開發儀表板","753 行","✅ 上線就緒（雙模式）","導覽、S1–S11、兩層引擎、健檢、定位"],
        ["simulator.html 整合策略沙盤 v4","610 行","🟡 可玩・待真實感調校","訓練：人格／家族／情緒／財務／競爭"],
        ["evaluator.html 坪效・開發試算","416 行","✅ 上線就緒（教學版）","容積帳→坪效→投報→敏感度，§162 驗證"],
        ["whitepaper.html 整合開發評估總論","324 行","✅ 完成","法規×估價×趨勢×人文最上位框架"],
        ["briefing.html ＋ .pptx 說明會簡報","480 行＋17 頁","✅ 範本完成","對住戶溝通（制式 8 段）"],
        ["整合開發評估總論 .md／.docx","—","✅ 完成","知識沉澱、開發人員閱讀"],
        ["方法論正典 ×2、spec ×5、產生器 ×2","—","✅ 完成","S1–S11／投報公式／工程留痕"],
    ],widths=[5.6,2.6,3.6,5.2])
    para(doc,"技術一致性：全部為單一自含 HTML（inline CSS＋原生 JS、零依賴、零建置），可離線、可上 GitHub Pages（已附 .nojekyll）。",9.5,MUTE)

    # 四、技術與品質
    heading(doc,"技術與品質評估 Technical & Quality",2,"04")
    heading(doc,"資安態勢（2026-06-22 掃描）",3)
    table(doc,["項目","結果"],[
        ["對外資料傳輸（fetch／beacon／analytics）","✅ 無（純前端、零外送、零追蹤）"],
        ["密鑰／token／secret","✅ 無"],["XSS 面（innerHTML 注入外部輸入）","✅ 無（皆內建字串模板）"],
        ["外連 target=_blank 之 noopener","✅ 21／21 全具備"],["localStorage","✅ 僅遊戲狀態與模式旗標，無 PII"],
        [".gitignore 鐵律","✅ 生效"],["真實案名／金額／人名洩漏","✅ 乾淨（已移除關係企業名稱）"],
        ["殘留（低風險）","🟡 Google Fonts 第三方請求；可改自託管字型達零外連"],
    ],widths=[8.0,9.0])
    heading(doc,"品質驗證（實測，非臆測）",3)
    bullets(doc,[
        ("靜態檢查","5 頁 JS 語法、HTML 標籤平衡、站內連結解析皆通過。"),
        ("沙盤實測","作者實際跑完一局並探針：人格差異化有效（說明會→情感+14.3、加碼→投資+16.7、承諾→保守+17.8）；家族溝通帶動全家族；財務隨決策／時間反應。"),
    ])

    # 五、法規查證
    heading(doc,"法規與市場查證 Regulatory & Market Verification",2,"05")
    para(doc,"以下為 2026-06 網路查證結果（來源見附錄 A）；本專案內容據此撰寫，非憑記憶。",9.5,MUTE)
    table(doc,["主題","查證結論（現行）"],[
        ["都更同意門檻 §37","自行劃定 3/4、政府劃定 1/2；面積逾 9/10 同意者人數不計"],
        ["都更容積獎勵 §65＋辦法","上限約基準容積 1.5 倍"],
        ["都更稅捐減免 §67","地價稅／房屋稅／土增稅／契稅減免"],
        ["危老容積獎勵 §6","上限 法定 1.3 倍 或 原建築 1.15 倍"],
        ["危老時程獎勵","已於民國 114（2025）年 5 月歸零落日 ← 2026 關鍵變動"],
        ["建築技術規則 §162","陽台≤10%／梯廳≤10%／合計≤15%／機電≤15%（逐層）；雨遮不計"],
        ["房地合一稅 2.0","≤2 年 45%、2–5 年 35%、5–10 年 20%、>10 年 15%；自住 6 年 10%（400 萬免稅額）"],
        ["囤房稅 2.0","2024/7 上路、2025/5 首徵；全國歸戶；非自住 2%–4.8%；自住 1%；影響約 36 萬戶"],
        ["平均地權條例（2023）","限制預售換約、私法人購屋許可制、檢舉獎金、重罰炒作 100 萬–5,000 萬"],
        ["建築能效 BERS／淨零","分級 1+~7；2024 版手冊 2025/7/1 實施；綠建築≥4 級；對接 2050 淨零"],
        ["2026 房市趨勢","價緩跌量盤整；餘屋 8–10 萬戶、完工 32–36 萬戶；營造年漲 2–3%；火力集中都更危老"],
    ],widths=[4.6,12.4])

    # 六、發現
    heading(doc,"關鍵發現 Key Findings",2,"06")
    para(doc,"強項：①內容經查證、標條號；②零依賴可上線；③雙模式服務外行與專業；④沙盤人格／家族／情緒機制已落地可驗證。",10,INK)
    heading(doc,"缺口（依影響排序，皆實測所得）",3)
    table(doc,["#","發現","證據","影響"],[
        ["G1","釘子戶不夠黏、難度偏易","第 3 回合清空反對、第 8 回合達 100% 支持","核心『人最難整合』感受不足"],
        ["G2","競爭建商為展示性","擺爛三場皆『時間到』敗，無一由對手達標","競爭機制無實質張力"],
        ["G3","財務單向、多結局形同虛設","ROI 僅隨時間下滑 32%→23%、無正向槓桿","財務缺可優化變數"],
        ["G4","回饋斷層","bandwagon 隱形、達標未催收尾、單卡近萬用","UX 與教學受限"],
    ],widths=[1.0,4.2,6.8,5.0])

    # 七、建議與展望
    heading(doc,"建議與執行展望 Recommendations & Outlook",2,"07")
    heading(doc,"即刻優先（低成本高回報）— 真實感調校包",3)
    bullets(doc,[
        ("釘子戶黏性","植入高抗性樞紐戶，僅特定高成本動作可撬動；達標後仍保留殘餘 holdout（權變強制）。"),
        ("競爭建商長牙","對手速度與『你的怠惰』連動，使『被搶先』成為真實敗局。"),
        ("回饋補強","bandwagon 即時提示、達標催收尾、情境逼出不同最優解。"),
    ])
    heading(doc,"中期 — RE-DCF-Tool 串接（修 G3）",3)
    para(doc,"RE-DCF-Tool 經查證為 Streamlit Python（L1–L6：§162 逐層容積查核、坪效、都更六科目投報、敏感度、三組黃金測試），無對外 API。三階段：")
    table(doc,["階段","做法","改 RE-DCF？","價值"],[
        ["Phase 0 公式同源","沙盤財務＋evaluator 採 RE-DCF 同一套 L2–L6 公式骨架；雙向 cross-link","零改","立即一致；ROI 取得正向槓桿"],
        ["Phase 1 參數交換","定義案件參數 JSON；深連結 st.query_params 或 JSON 匯入匯出","小改","一鍵送 RE-DCF 精算"],
        ["Phase 2 計算 API","RE-DCF 加 FastAPI／/api 回 KPI／IRR JSON","較大（後端＋CORS）","啟用案件生成器、自動化"],
    ],widths=[3.0,7.0,2.6,4.4])
    callout(doc,"資安鐵則","RE-DCF 真實案夾已 gitignore；任何串接一律『使用者輸入、不落地公庫』。建議追加『真實數字本地模式』（僅存瀏覽器、永不上傳）供內訓安全使用。","warn")
    heading(doc,"長期 — V4→V7 路線圖",3)
    bullets(doc,[
        ("V4（進行中）","財務系統 ✅、競爭建商（待長牙）、法規風險擴充 ✅。"),
        ("V5","RE-DCF 串接、案件生成器（使用者輸入生成通用劇本）。"),
        ("V6","AI 地主代理人（需 LLM 後端，靜態頁無法承載，須改架構）。"),
        ("V7","都市更新數位決策平台（知識庫＋計算工具＋AI 地主整合）。"),
    ])
    heading(doc,"上線與治理 Outlook",3)
    bullets(doc,[
        ("上線","GitHub Settings → Pages → main /（root）啟用，即得 jeremy0819.github.io/Urban-Renewal/。"),
        ("建議 KPI","開發部試用回收、Pages 流量、訓練覆盤產出數。"),
        ("風險","①AI 地主需後端；②真實數字一律本地、嚴守不入庫；③Google Fonts 可自託管消除外連。"),
    ])

    # 附錄
    doc.add_page_break()
    heading(doc,"附錄 A — 查證來源（查證日 2026-06-22）",2)
    for s in [
        "房地合一稅 2.0：今周刊；財政部北區國稅局",
        "囤房稅 2.0（全國歸戶、2%–4.8%、2025/5 首徵）：財政部房屋稅 2.0；商業周刊；關鍵評論網",
        "建築能效 BERS／淨零：內政部建築研究所（2024 版手冊 114/7/1 實施）；澄毓綠建築；CSRone",
        "都更條例 §37／§65／§67：全國法規資料庫；都市更新建築容積獎勵辦法",
        "危老時程獎勵落日：聯合新聞網；安信建經",
        "建築技術規則 §162：全國法規資料庫；房感",
        "平均地權條例 2023 修正：內政部",
        "2026 房市趨勢：遠雄房地產（劉佩真）；鉅亨網",
        "RE-DCF-Tool 技術內容：GitHub README（jeremy0819/RE-DCF-Tool）",
    ]:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(1); run(p,s,9.2,BODY)
    heading(doc,"附錄 B — Repo 稽核摘要",2)
    bullets(doc,[
        ("提交數","18（最新 5653e22）。"),
        ("程式碼","5 個 HTML 共 2,583 行（index 753／simulator 610／evaluator 416／whitepaper 324／briefing 480）。"),
        ("文件","方法論 ×3、spec ×5、Word ×1、PPT ×1、產生器 ×2。"),
        ("洩漏掃描","✅ 無真實案名／金額／人名／密鑰。"),
    ])
    para(doc,"免責：本報告之法規／稅率／市場數據以查證日之公開資料為準，實務適用須以現行條文、主管機關核定與專業查證為據；本專案為通用方法論與教學工具，非正式財務／法律意見。",8.5,MUTE,before=10,line=1.35)

    out="Urban-Renewal-狀態報告-2026-06.docx"; doc.save(out); print("✓ 已輸出",out); return out

if __name__=="__main__": build()
