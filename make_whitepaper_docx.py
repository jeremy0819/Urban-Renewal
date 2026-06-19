#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
都更／危老 整合開發評估總論（2026）→ Word (.docx) 產生器
供開發人員閱讀的長文白皮書，內容同 docs/整合開發評估總論-2026.md 與 whitepaper.html。
品牌：jieceng 暖白 × 炭黑 × 翡翠綠（Noto Serif TC 標題 / 微軟正黑內文）。
通用方法論，不含真實案件資料；法規標條號、行情標查證日 2026-06-19。

用法：  python3 make_whitepaper_docx.py
輸出：  整合開發評估總論-2026.docx
相依：  pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EMER  = RGBColor(0x06, 0x4E, 0x3B)
EMERL = RGBColor(0x10, 0x98, 0x71)
INK   = RGBColor(0x1C, 0x1A, 0x17)
BODY  = RGBColor(0x44, 0x40, 0x3C)
MUTE  = RGBColor(0x8A, 0x83, 0x7C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
WARM  = "F1EEE6"     # 淺暖底（cell shading）
MINT  = "E9F7F0"     # 淺薄荷（TL;DR）
WARNB = "FBF3E6"     # 淺琥珀（警示）
SERIF = "Noto Serif TC"
SANS  = "Microsoft JhengHei"


def _ea(run, name):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), name)
    rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name)


def shade(cell, hexcolor):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def run(p, text, size=10.5, color=BODY, bold=False, font=SANS):
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = color; r.font.name = font; _ea(r, font)
    return r


def para(doc, text="", size=10.5, color=BODY, bold=False, font=SANS,
         align=None, before=2, after=4, line=1.32):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if align: p.alignment = align
    if text: run(p, text, size, color, bold, font)
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        if isinstance(it, tuple):     # (粗體前綴, 其餘)
            run(p, it[0] + '：', size, INK, bold=True); run(p, it[1], size, BODY)
        else:
            run(p, it, size, BODY)


def heading(doc, text, level, kicker=None):
    if kicker:
        k = doc.add_paragraph(); k.paragraph_format.space_before = Pt(14); k.paragraph_format.space_after = Pt(0)
        run(k, kicker.upper(), 8.5, EMERL, bold=True, font=SANS)
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(2 if kicker else 12); pf.space_after = Pt(5)
    sizes = {1: 22, 2: 16, 3: 12.5}
    run(p, text, sizes.get(level, 12), EMER if level <= 2 else INK, bold=True, font=SERIF)
    return p


def callout(doc, label, text, kind="emer"):
    fill = {"emer": MINT, "warn": WARNB}.get(kind, MINT)
    bar = {"emer": EMER, "warn": RGBColor(0xC2, 0x70, 0x1C)}.get(kind, EMER)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); shade(c, fill)
    c.paragraphs[0].paragraph_format.space_after = Pt(1)
    run(c.paragraphs[0], label, 8.5, bar, bold=True)
    p2 = c.add_paragraph(); p2.paragraph_format.space_before = Pt(1)
    run(p2, text, 10, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tldr(doc, text):
    callout(doc, "TL;DR", text, "emer")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "0A5D46")
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(1)
        hdr[i].paragraphs[0].paragraph_format.space_before = Pt(1)
        run(hdr[i].paragraphs[0], h, 9, WHITE, bold=True)
    for r_i, rowdata in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(rowdata):
            if r_i % 2 == 1: shade(cells[i], WARM)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(1)
            cells[i].paragraphs[0].paragraph_format.space_before = Pt(1)
            run(cells[i].paragraphs[0], val, 9, BODY)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================ 建構 ============================
def build():
    doc = Document()
    st = doc.styles['Normal']; st.font.name = SANS; st.font.size = Pt(10.5)
    _ea(st.element.get_or_add_rPr().__class__ and st.element, SANS) if False else None
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.2)

    # ---- 封面 ----
    para(doc, before=40)
    para(doc, "WHITE PAPER · 整合人的視角", 10, EMERL, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, "都更／危老\n整合開發評估總論", 30, EMER, bold=True, font=SERIF)
    para(doc, "以不動產研析為脊椎，縫合法制、財務與人本", 14, BODY, bold=True, before=4, after=18)
    para(doc, "都更不是工程問題，是「法制 × 財務 × 人」的三體問題。本白皮書把都市更新"
              "法規、都市計畫、建築估價、稅制與 2026 房地產趨勢疊成一張判斷圖，並以"
              "「世界唯一的整合人」視角，提出一套最獨特的開發評估框架。", 11, BODY, after=22, line=1.5)
    table(doc, ["項目", "內容"], [
        ["查證日", "2026-06-19"],
        ["定位", "方法論最上位思維框架（統攝 S1–S11 與投報兩層）"],
        ["邊界", "通用方法論・不含任何真實案件資料"],
        ["閱讀對象", "開發人員／整合團隊／評估分析"],
    ], widths=[3.5, 13])

    doc.add_page_break()

    # ---- 目錄 ----
    heading(doc, "目錄", 2)
    for t_ in ["序章　為什麼需要「整合人」", "第一部　法規宇宙：把所有規則疊成一張地圖",
               "第二部　都市計畫上位整合", "第三部　不動產研析作為分析脊椎",
               "第四部　2026 趨勢透鏡", "第五部　數據邏輯 × 人文思考的縫合",
               "第六部　整合人評估框架（最獨特的角度）", "第七部　報告格式（脈絡）",
               "結語　整合人的三條鐵律", "參考來源（查證日 2026-06-19）"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        run(p, t_, 10.5, INK)

    doc.add_page_break()

    # ---- 序章 ----
    heading(doc, "為什麼需要「整合人」", 2, "序章")
    para(doc, "建築師談量體、估價師談價值、律師談條款、代銷談去化、地主談記憶與恐懼——"
              "每個專業都對，但每個只拿著一面透鏡。案子卡住，幾乎從不是「沒有人會算」，"
              "而是沒有人把所有透鏡疊在同一張圖上。")
    para(doc, "整合人的定位：不是說服者，是判斷者；不是中介，是縫合者。第一件事不是推案，"
              "是先查清楚「能不能做、划不划算、整不整得起來」三件事的交集。")
    callout(doc, "一句話心法", "把餅做大是數學，把人留下是人文；整合人要同時是會計師與人類學家。")

    # ---- 第一部 ----
    heading(doc, "法規宇宙：把所有規則疊成一張地圖", 2, "第一部")
    tldr(doc, "核心三法決定路徑與門檻，容積四源決定量體真相，估價與權變決定怎麼分，"
              "稅制決定誰負擔——而真正讓案子翻車的，常是「你沒想過的那一條」。")
    heading(doc, "核心三法・同意門檻", 3)
    table(doc, ["法源", "角色", "同意門檻（依現行條文）"], [
        ["都市更新條例", "多數決整合主幹", "事業概要§22 >1/2；事業計畫§37：第7條地區 >1/2；其餘（含自行劃定）>3/4；面積 >9/10 者人數不計"],
        ["都市計畫法", "上位：分區・容積・都審", "決定「可建範圍」的法定基礎"],
        ["危老條例", "小基地快速通道", "全體（100%）同意，無須劃定更新單元"],
    ], widths=[3.2, 4.2, 9.1])
    callout(doc, "2026 關鍵變動", "危老時程獎勵已落日歸零（8→6→4→2→1%，至民國114／2025年5月歸零）。"
            "急迫性消退，路徑天平重新傾向都更多數決。", "warn")
    heading(doc, "容積的真相：量體只有一個來源", 3)
    table(doc, ["來源", "重點"], [
        ["法定容積", "容積率 × 土地面積；FA 用使照基地面積，非謄本"],
        ["都更容積獎勵（辦法・條例§65）", "上限法定 1.5 倍，或原建築容積 + 0.3 倍基準；容獎特快車擬直給1.5倍、捐社宅疊加至2倍"],
        ["危老容積獎勵（條例§6）", "上限法定 1.3 倍 或 原建築 1.15 倍；時程＋規模 ≤10%（時程已歸零）"],
        ["§162 免計容積", "逐層核算（機電／梯廳／陽台）；總量法放大近 10 倍"],
        ["容積移轉 TDR", "古蹟／公設地容積移入，受移入上限管制"],
    ], widths=[5.0, 11.3])
    heading(doc, "你可能沒想過的那幾條（盲區清單）", 3)
    para(doc, "真正讓案子翻車的，常不是核心三法，而是這些邊界法規——整合人當前置風險逐項排雷：")
    table(doc, ["法源", "為什麼會翻車"], [
        ["平均地權條例（2023/7）", "限制預售換約、私法人購屋許可制（5年內不得移轉）、重罰炒作 → 改變買方結構與去化"],
        ["公寓大廈管理條例", "交屋後管委會、規約、公設點交；更新前約定專用影響拆除同意"],
        ["信託法・不動產開發信託", "續建機制與資金專戶的法制基礎——制度信用取代公司信用"],
        ["文化資產保存法", "老屋可能被列冊／指定，影響可拆性與容積（亦可換容積）"],
        ["地質法・環評法", "活動斷層／地質敏感區、一定規模需環評 → 可行性與時程風險"],
        ["民法（共有・優先承買）", "持分、繼承未分割、優先承買權——權屬整合的硬骨頭"],
        ["銀行法§72-2・土建融管制", "放款水位與央行管制，直接決定資金到不到位（2026 重中之重）"],
    ], widths=[4.4, 11.9])

    # ---- 第二部 ----
    heading(doc, "都市計畫上位整合：先讀計畫，再談容積", 2, "第二部")
    tldr(doc, "容積率不是憑空的數字，它從上位計畫一路長下來；不讀都市計畫就算容積，等於在沙上蓋樓。")
    bullets(doc, [
        ("上位計畫 → 分區管制 → 容積率／建蔽率", "先確認主要／細部計畫、退縮、開放空間、高度管制"),
        ("都市設計審議", "一定規模須過都審，影響量體、外觀與時程"),
        ("TOD／都市再生／策略地區", "常帶額外容積，是爭取量體的槓桿"),
        ("容積移轉（TDR）", "受移入比例與區位上限管制，須與獎勵一併納入「容積帳」"),
    ])
    callout(doc, "整合人動作", "在 S2 之前先做「都市計畫體檢表」：上位計畫、分區、容積率、都審門檻、"
            "再生／TOD 機會、容移可行性逐項查清，再進容積試算。")

    # ---- 第三部 ----
    heading(doc, "不動產研析作為分析脊椎", 2, "第三部")
    tldr(doc, "研析不是「算一個數字」，是建立一條可被挑戰的推理鏈——從最有效使用，到估價，到現金流，到風險。")
    heading(doc, "最有效使用分析（HBU）", 3)
    para(doc, "四道篩網，缺一不可：① 法律允許 → ② 實質可能 → ③ 財務可行 → ④ 最大效益。"
              "整合人先問「該蓋什麼」，而非「能蓋多少」。")
    heading(doc, "估價三方法・權變的更新前後權利價值", 3)
    table(doc, ["方法", "在都更的落點"], [
        ["比較法", "更新前權利價值、銷售單價假設（S3–S4）"],
        ["收益法", "店舖、收益型物業、整體財務（S4）"],
        ["成本法", "更新後評估、營建成本基準（S4・S7）"],
    ], widths=[3.5, 12.8])
    para(doc, "權變估價：更新前權利價值由實施者委託 ≥3 家估價者查估，三家差距 20% 以內擇一並說明；"
              "未建築者以素地價值推估。估價是都更最常見的爭議引爆點——基準與選任須攤在陽光下。")
    heading(doc, "財務研析：敏感度三軸", 3)
    para(doc, "售價、營建單價、時程——2026 三者皆逆風（見第四部）。報告不是給一個數字，"
              "是給一條標明假設、可被反駁的推理鏈。")

    # ---- 第四部 ----
    heading(doc, "2026 趨勢透鏡：把時間軸納入評估", 2, "第四部")
    tldr(doc, "（查證 2026-06-19）高利率、餘屋賣壓、營造成本續漲、人口轉向、政策降溫、淨零碳上路——"
              "六股風同吹，且多數對都更危老是逆風中的順風（資金被擠進更新題材）。")
    table(doc, ["透鏡", "2026 現況", "回寫評估"], [
        ["資金・利率", "價緩跌量盤整；央行轉精準控管・用途引導，利率微降；土建融未鬆", "折現率保守；素地難求＋限貸 → 火力集中都更危老"],
        ["供給・餘屋", "餘屋約 8–10 萬戶；2026–28 完工 32–36 萬戶，賣壓沉重", "去化速度下修，回寫現金流與敏感度"],
        ["營造成本", "年漲 2–3%；台北 RC ≈每坪28萬、鋼構36萬、其他都會18–20萬", "用總樓地板坪、升幅入模；共負上升→分回壓縮"],
        ["人口結構", "高齡化、少子化、家戶縮小", "產品轉小宅・適老・無障礙"],
        ["政策降溫", "平均地權・囤房稅2.0（全國歸戶2%–4.8%）・房地合一2.0", "買方轉自用・長期持有・法人退場"],
        ["淨零・能效", "建築能效 BERS（2025/7實施，分級1+~7）；2040前50%既有達能效一級", "危老都更＝國家級減碳槓桿；能效＋綠建築疊加容獎、對接 ESG 資金"],
    ], widths=[2.6, 7.2, 6.5])
    callout(doc, "趨勢的整合啟示", "2026 是「逆風中的結構性順風」——市場降溫，但限貸與淨零政策把資金與題材"
            "擠向都更危老。整合人要把六股風逐一回寫進敏感度，而非只喊「都更有前景」。")

    # ---- 第五部 ----
    heading(doc, "數據邏輯 × 人文思考的縫合", 2, "第五部")
    tldr(doc, "數據決定可行性（feasible），人文決定可成性（achievable）。能算出來的案子未必整合得起來；差別在人。")
    bullets(doc, [
        ("把信任當資本", "整合的真正貨幣是信任。制度信用 > 公司信用——信託、續建機制、連帶保證、單一數源"),
        ("釘子戶不是貪婪，是恐懼", "多數「拿翹」其實是資訊不對稱下的恐懼型／理性質疑型。專業質疑是養分，不是阻礙"),
        ("競合而非零和", "整合成功才有餅（正和），破局則餅歸零（雙輸）。合作把餅做大、競爭分餅"),
        ("程序正義的記憶", "聽證、公展、估價選任每一步留痕，是制度成本，也是社會信任的修復"),
    ])
    callout(doc, "縫合準則", "先用數據證明「值得做」，再用人文確保「做得成」；報告若只有前者，是半份報告。")

    # ---- 第六部 ----
    heading(doc, "整合人評估框架（最獨特的角度）", 2, "第六部")
    para(doc, "假設世界上只有一個都更整合人，他的簽名工具是「三層透鏡疊圖 × 短板診斷」。", bold=False)
    table(doc, ["能不能做", "划不划算", "整不整得起來"], [
        ["法制層", "財務層", "人本層"],
        ["容積・門檻・路徑", "投報・估價・稅", "信任・賽局・程序"],
    ], widths=[5.4, 5.4, 5.5])
    para(doc, "整合人坐落於三者交集。", 10, MUTE, align=WD_ALIGN_PARAGRAPH.CENTER)
    heading(doc, "短板決定成敗（不是加總，是取最小）", 3)
    para(doc, "傳統評估把各項加權相加；整合人反過來——三層各打一分，看最低的那一層。容積再漂亮（90）、"
              "投報再香（85），只要人本層是 20（信任崩、釘子戶僵局），案子就是 20 分。"
              "整合人先找短板，再決定要不要、以及如何補。")
    heading(doc, "五個診斷問題", 3)
    bullets(doc, [
        ("法制", "路徑選對了嗎？容積帳每分獎勵都對得上法源上限嗎？（§37／§65／§162）"),
        ("財務", "在售價／營建／時程三逆風下，保守情境還活著嗎？"),
        ("估價", "更新前權利價值的基準與選任透明嗎？三家收斂嗎？"),
        ("人本", "同意率是真實的，還是偏好偽裝堆出的紙牌屋？短板戶是恐懼型還是拿翹型？"),
        ("制度信用", "續建機制、信託、擔保、單一數源到位了嗎？"),
    ])
    callout(doc, "整合人的自我定位", "當判斷者，不當說服者；先查 corporate substance，再談其他。"
            "同時看「現在的數字」與「會怎麼演變的人」。")

    # ---- 第七部 ----
    heading(doc, "報告格式（脈絡）：每份評估都這樣寫", 2, "第七部")
    para(doc, "報告不是資料堆疊，是一條有起承轉合的推理脈絡。從「這塊地的故事」開始，到「整合人的判斷」結束。")
    table(doc, ["段", "標題", "脈絡"], [
        ["0", "一頁判斷", "三層分數＋短板＋去留建議，先講結論"],
        ["1", "基地的故事", "區位、權屬、屋齡、紋理（人文起手）"],
        ["2", "法制體檢", "路徑、上位計畫、容積帳、§162 逐層"],
        ["3", "不動產研析", "最有效使用、估價三方法、權利價值"],
        ["4", "財務與敏感度", "投報兩層、共負六科目、三逆風情境"],
        ["5", "趨勢回寫", "利率／餘屋／營造／人口／政策／淨零"],
        ["6", "人本與賽局", "同意結構、短板戶辨型、制度信用解方"],
        ["7", "整合人判斷", "短板、條件、時程賽局下的最適路徑"],
    ], widths=[1.4, 4.0, 10.9])
    callout(doc, "寫法準則", "每一節都回答「所以呢？」——資料之後一定接判斷；判斷之後一定標假設與來源。")

    # ---- 結語 ----
    heading(doc, "整合人的三條鐵律", 2, "結語")
    bullets(doc, [
        ("三層疊圖、短板決定", "能不能做 × 划不划算 × 整不整得起來，看最低的那一層"),
        ("數據證明值得、人文確保做成", "兩者缺一，都是半份報告"),
        ("制度信用 > 公司信用", "用信託、續建、擔保、單一數源、透明估價，把信任變成可計量、可交付的資本"),
    ])
    para(doc, "都更的終局不是蓋好一棟樓，是讓一群人願意一起把家交給未來。整合人是那個同時握著計算機與同理心的人。",
         10.5, EMER, bold=True)

    # ---- 來源 ----
    heading(doc, "參考來源（查證日 2026-06-19）", 2)
    for s in [
        "全國法規資料庫：都市更新條例§37、§67；不動產估價技術規則；都市更新權利變換實施辦法",
        "內政部：都市更新建築容積獎勵辦法；建築能效評估手冊 BERS（114/7/1 實施）；平均地權條例修正",
        "危老時程獎勵落日：聯合新聞網、安信建經",
        "稅制：房地合一稅 2.0（財政部 QA）；囤房稅 2.0（全國歸戶）",
        "淨零建築：2050 淨零・建築能效規章（CSRone）",
        "2026 房市趨勢：遠雄房地產（劉佩真觀點）、鉅亨網",
    ]:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
        run(p, s, 9.5, BODY)
    para(doc, "免責：本白皮書為通用方法論與趨勢研析，非正式財務／法律意見；所有法規數字、稅率、"
              "獎勵上限與市場數據均以現行條文、主管機關核定及查證日後最新資料為準，個案適用須由"
              "法規與估價專業逐項查證。不含任何真實案件資料。版本 1.0｜2026-06。",
         8.5, MUTE, before=10, line=1.4)

    out = "整合開發評估總論-2026.docx"
    doc.save(out)
    print(f"✓ 已輸出 {out}")
    return out


if __name__ == "__main__":
    build()
